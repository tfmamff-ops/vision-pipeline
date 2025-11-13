import io
import json
import logging
import uuid

import azure.functions as func
import cv2
import numpy as np
import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

from shared_code.storage_util import download_bytes, upload_bytes

TEMPLATES_CONTAINER = "report-templates"
TEMPLATE_ACCEPT = "accept.docx"
TEMPLATE_REJECT = "reject.docx"
CLOUDMERSIVE_URL = "https://api.cloudmersive.com/convert/docx/to/pdf"
CLOUDMERSIVE_API_KEY = "d2e4f77e-0140-4072-9390-1ffcfbe2b1e9"
MIME_JSON = "application/json"


# ----------------------------------------------------------------------
# REPORT CONTENT PREPARATION
# ----------------------------------------------------------------------


def get_report_replacements_and_image_paths(instance_id, user_comment, accepted):
    """
    Build the replacements dictionary and image path metadata used to fill the DOCX template.
    """
    replacements = {
        "{{instance_id}}": "INS-20251010-001",
        "{{created_date}}": "2025/10/10",
        "{{created_time}}": "13:00:45",
        "{{requested_by_user_id}}": "user123",
        "{{requested_by_user_name}}": "Álvaro Morales",
        "{{requested_by_user_email}}": "alvaro.m@company.com",
        "{{requested_by_user_role}}": "Operator",
        "{{client_app_version}}": "1.5.2",
        "{{expected_prod_code}}": "PROD-ABC-456",
        "{{expected_prod_desc}}": "Paracetamol 500mg - Tablets",
        "{{expected_lot}}": "LOT-2025-09-001",
        "{{validation_lot_ok}}": "✔",
        "{{expected_exp_date}}": "2027/12/31",
        "{{validation_exp_date_ok}}": "✔",
        "{{expected_pack_date}}": "2025/09/20",
        "{{validation_pack_date_ok}}": "✘",
        "{{validation_barcode_detected_ok}}": "✘",
        "{{validation_barcode_legible_ok}}": "✔",
        "{{barcode_payload_decoded_value}}": "GS1-98765432101234",
        "{{barcode_payload_barcode_symbology}}": "DataMatrix",
        "{{input_container}}": "cont-in-2025",
        "{{input_blob_name}}": "input_001.jpg",
        "{{processed_image_container}}": "cont-proc-2025",
        "{{processed_image_blob_name}}": "processed_001.jpg",
        "{{ocr_overlay_container}}": "cont-ocr-2025",
        "{{ocr_overlay_blob_name}}": "ocr_overlay_001.png",
        "{{barcode_overlay_container}}": "cont-bar-2025",
        "{{barcode_overlay_blob_name}}": "barcode_overlay_001.png",
        "{{barcode_roi_container}}": "cont-roi-2025",
        "{{barcode_roi_blob_name}}": "barcode_roi_001.png",
        "{{VALOR_AND}}": "✔",
        "{{validation_barcode_ok}}": "✔",
        "{{validation_summary}}": "✘",
        "{{user_comment}}": "Product successfully verified. Dates are visible and legible.",
        "{{report_container}}": "cont-report-2025",
        "{{report_blob_name}}": "report_001.docx",
    }

    image_paths = {
        "input_image": {
            "container": "input",
            "blobName": "uploads/0a0643fc-fd95-480f-94cc-459f21d03aeb.jpg",
            "resizePercentage": 40,
        },
        "processed_image": {
            "container": "output",
            "blobName": "final/ocr/processed/03de1e46-ede1-4354-a890-69b550c08c33.png",
            "resizePercentage": 40,
        },
        "ocr_overlay_image": {
            "container": "output",
            "blobName": "final/ocr/overlay/04f75521-4400-4d79-8e30-ded2952200a9.png",
            "resizePercentage": 40,
        },
        "barcode_overlay_image": {
            "container": "output",
            "blobName": "final/barcode/overlay/02d7cd2d-5913-4778-a7ea-b0093bb75f45.png",
            "resizePercentage": 40,
        },
        "barcode_roi_image": {
            "container": "output",
            "blobName": "final/barcode/roi/02d7cd2d-5913-4778-a7ea-b0093bb75f45.png",
            "resizePercentage": 40,
        },
    }

    return replacements, image_paths


# ----------------------------------------------------------------------
# DOCX → PDF CONVERSION (CLOUDMERSIVE)
# ----------------------------------------------------------------------


def convert_docx_to_pdf_cloudmersive(docx_bytes: io.BytesIO, api_key: str) -> bytes | None:
    """
    Send a DOCX in memory to Cloudmersive and return the resulting PDF bytes,
    or None if the conversion fails.
    """
    payload = docx_bytes.getvalue()
    logging.info("[generate_report] Cloudmersive upload DOCX size: %d bytes", len(payload))

    headers = {"Apikey": api_key}
    files = {
        "inputFile": (
            "input.docx",
            payload,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }

    try:
        response = requests.post(CLOUDMERSIVE_URL, headers=headers, files=files, timeout=60)
    except requests.RequestException as exc:
        logging.error("[generate_report] Cloudmersive network error: %s", exc)
        return None

    if response.status_code == 200:
        return response.content

    logging.error(
        "[generate_report] Cloudmersive error %s. Response: %s",
        response.status_code,
        response.text[:300],
    )
    return None


# ----------------------------------------------------------------------
# DOCX REPORT GENERATION
# ----------------------------------------------------------------------


def add_colored_text(paragraph: Paragraph, text: str) -> None:
    """
    Write text into a paragraph and color ✔ in green and ✘ in red.
    Other characters keep the default color.
    """
    for ch in text:
        run = paragraph.add_run(ch)
        if ch == "✔":
            run.font.color.rgb = RGBColor(0, 150, 0)
        elif ch == "✘":
            run.font.color.rgb = RGBColor(200, 0, 0)


def clear_paragraph_runs(paragraph: Paragraph) -> None:
    """Remove all runs from a paragraph."""
    for run in reversed(paragraph.runs):
        paragraph._element.remove(run._element)


def try_insert_image(paragraph: Paragraph, full_text: str, image_paths: dict) -> bool:
    """
    Look for an image placeholder in the paragraph and, if found,
    download and insert the corresponding image.
    """
    for img_placeholder, img_info in image_paths.items():
        token = f"{{{{{img_placeholder}}}}}"
        if token not in full_text:
            continue

        clear_paragraph_runs(paragraph)

        container = img_info.get("container")
        blob_name = img_info.get("blobName")
        resize_pct = img_info.get("resizePercentage", 40)

        final_img_source = get_image(container, blob_name, resize_percentage=resize_pct)

        if final_img_source:
            run = paragraph.add_run()
            run.add_picture(final_img_source)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return True

        logging.warning(
            "[generate_report] No valid image for placeholder '%s' (container=%r, blobName=%r)",
            img_placeholder,
            container,
            blob_name,
        )
        return False

    return False


def apply_text_replacements(paragraph: Paragraph, replacements: dict) -> None:
    """
    Replace all placeholder tokens in a paragraph with their corresponding values.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text

    for placeholder, value in replacements.items():
        new_text = new_text.replace(placeholder, str(value))

    if new_text != full_text:
        clear_paragraph_runs(paragraph)
        add_colored_text(paragraph, new_text)


def iter_paragraphs(element) -> list[Paragraph]:
    """
    Return a flat list of Paragraph objects from either a single paragraph or a table.
    """
    if isinstance(element, Paragraph):
        return [element]

    if isinstance(element, Table):
        result: list[Paragraph] = []
        for row in element.rows:
            for cell in row.cells:
                result.extend(cell.paragraphs)
        return result

    return []


def replace_in_element(element, replacements: dict, image_paths: dict) -> None:
    """
    Apply image replacement and text replacement on all paragraphs inside an element.
    """
    for paragraph in iter_paragraphs(element):
        full_text = "".join(run.text for run in paragraph.runs)

        image_inserted = try_insert_image(paragraph, full_text, image_paths)
        if not image_inserted:
            apply_text_replacements(paragraph, replacements)


def generate_verification_report_bytes(
    template: bytes,
    replacements: dict,
    image_paths: dict,
) -> io.BytesIO | None:
    """
    Load a DOCX template from bytes, apply text and image placeholders,
    and return a new DOCX as BytesIO. Returns None if the template cannot be loaded.
    """
    try:
        buf = io.BytesIO(template)
        document = Document(buf)
    except Exception:
        logging.exception("[generate_report] Error loading DOCX template")
        return None

    # Process paragraphs
    for paragraph in document.paragraphs:
        replace_in_element(paragraph, replacements, image_paths)

    # Process tables
    for table in document.tables:
        replace_in_element(table, replacements, image_paths)

    out_stream = io.BytesIO()
    document.save(out_stream)
    out_stream.seek(0)
    return out_stream


# ----------------------------------------------------------------------
# IMAGE SUPPORT CODE
# ----------------------------------------------------------------------


def resize_by_percentage(img: np.ndarray, percentage: int) -> np.ndarray | None:
    """
    Scale an OpenCV image (ndarray) by a percentage.
    Returns a resized ndarray or None if validation or resize fails.
    """
    if img is None:
        logging.error("[resize_by_percentage] input image is None")
        return None

    if percentage <= 0:
        logging.error("[resize_by_percentage] percentage must be > 0")
        return None

    scale = percentage / 100.0
    new_w = int(img.shape[1] * scale)
    new_h = int(img.shape[0] * scale)

    if new_w == 0 or new_h == 0:
        logging.error("[resize_by_percentage] resulting size is zero (%d×%d)", new_w, new_h)
        return None

    try:
        resized = cv2.resize(img, (new_w, new_h), interpolation=cv2.INTER_AREA)
    except Exception as exc:
        logging.error("[resize_by_percentage] cv2.resize failed: %s", exc)
        return None

    return resized


def get_image(container: str, blob_name: str, resize_percentage: int = 50) -> io.BytesIO | None:
    """
    Download an image from Blob Storage, validate and resize it,
    and return the final PNG as a BytesIO stream.
    """
    try:
        img_bytes = download_bytes(container, blob_name)
    except Exception as exc:
        logging.error("[generate_report] error downloading image: %s", exc)
        return None

    if not img_bytes:
        logging.error("[generate_report] image blob is empty")
        return None

    arr = np.frombuffer(img_bytes, dtype=np.uint8)
    img = cv2.imdecode(arr, cv2.IMREAD_COLOR)

    if img is None:
        logging.error("[generate_report] image is NOT a valid image (imdecode returned None)")
        return None

    logging.info(
        "[generate_report] image OK: shape=%s dtype=%s",
        img.shape,
        img.dtype,
    )

    resized = resize_by_percentage(img, resize_percentage)
    if resized is None:
        logging.error("[generate_report] resize_by_percentage returned None")
        return None

    success, encoded_png = cv2.imencode(".png", resized)
    if not success:
        logging.error("[generate_report] failed to encode resized image to PNG")
        return None

    buf = io.BytesIO(encoded_png.tobytes())
    buf.seek(0)
    return buf


# ----------------------------------------------------------------------
# HTTP ENTRYPOINT (AZURE FUNCTION)
# ----------------------------------------------------------------------


def main(req: func.HttpRequest) -> func.HttpResponse:
    """
    HTTP-triggered entrypoint that generates a PDF report from DOCX templates
    stored in Blob Storage and returns the Blob reference of the final PDF.
    """
    try:
        payload = req.get_json()

        # Log the full received JSON body for auditing/troubleshooting purposes
        logging.info("[generate_report] JSON received (raw): %s", req.get_body().decode("utf-8"))
        logging.info(
            "[generate_report] Parsed JSON (payload): %s",
            json.dumps(payload, indent=2, ensure_ascii=False),
        )

        instance_id = payload.get("instanceId")
        user_comment = payload.get("userComment")
        accepted = payload.get("accepted")
    except Exception:
        logging.exception("[generate_report] Error processing JSON - returning 400")
        return func.HttpResponse(
            json.dumps(
                {
                    "ok": False,
                    "error": {"code": "invalid_json", "message": "Invalid JSON in request body"},
                }
            ),
            status_code=400,
            mimetype=MIME_JSON,
        )

    try:
        template_bytes = download_bytes(
            TEMPLATES_CONTAINER,
            TEMPLATE_ACCEPT if accepted else TEMPLATE_REJECT,
        )
    except Exception as exc:
        logging.exception("[generate_report] Failed to download template: %s", exc)
        return func.HttpResponse(
            json.dumps(
                {
                    "ok": False,
                    "error": {
                        "code": "missing_template",
                        "message": "Template could not be downloaded",
                    },
                }
            ),
            status_code=404,
            mimetype=MIME_JSON,
        )

    # Build the data used to fill the DOCX template
    replacements, image_paths = get_report_replacements_and_image_paths(
        instance_id,
        user_comment,
        accepted,
    )
    docx_stream = generate_verification_report_bytes(template_bytes, replacements, image_paths)

    if docx_stream is None:
        logging.error("[generate_report] DOCX generation failed")
        return func.HttpResponse(
            json.dumps(
                {
                    "ok": False,
                    "error": {
                        "code": "docx_generation_failed",
                        "message": "Could not generate DOCX report",
                    },
                }
            ),
            status_code=500,
            mimetype=MIME_JSON,
        )

    pdf_bytes = convert_docx_to_pdf_cloudmersive(docx_stream, CLOUDMERSIVE_API_KEY)

    if not pdf_bytes:
        logging.error("[generate_report] DOCX to PDF conversion failed")
        return func.HttpResponse(
            json.dumps(
                {
                    "ok": False,
                    "error": {
                        "code": "conversion_failed",
                        "message": "Could not convert DOCX to PDF",
                    },
                }
            ),
            status_code=502,
            mimetype=MIME_JSON,
        )

    out_blob_name = f"final/report/{uuid.uuid4()}.pdf"

    try:
        upload_bytes("output", out_blob_name, pdf_bytes, content_type="application/pdf")
        logging.info("[generate_report] Uploaded report as %s/%s", "output", out_blob_name)
    except Exception as exc:
        logging.exception("[generate_report] Failed to upload report: %s", exc)
        return func.HttpResponse(
            json.dumps(
                {
                    "ok": False,
                    "error": {
                        "code": "upload_failed",
                        "message": "Could not upload PDF to Blob Storage",
                    },
                }
            ),
            status_code=500,
            mimetype=MIME_JSON,
        )

    return func.HttpResponse(
        json.dumps(
            {
                "reportBlob": {
                    "container": "output",
                    "blobName": out_blob_name,
                }
            }
        ),
        status_code=200,
        mimetype=MIME_JSON,
    )
