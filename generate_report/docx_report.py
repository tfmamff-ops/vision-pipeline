import io
import logging

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

from generate_report.replacements import CHECK_MARK, CROSS_MARK

from .image_utils import get_image, get_unavailable_image

logger = logging.getLogger(__name__)

GREEN = RGBColor(0, 150, 0)
RED = RGBColor(200, 0, 0)


def add_colored_text(paragraph: Paragraph, text: str) -> None:
    """Write text into a paragraph and color ✔ in green and ✘ in red."""

    for ch in text:
        run = paragraph.add_run(ch)
        if ch == CHECK_MARK:
            run.font.color.rgb = GREEN
        elif ch == CROSS_MARK:
            run.font.color.rgb = RED


def clear_paragraph_runs(paragraph: Paragraph) -> None:
    """Remove all runs from a paragraph."""
    for run in reversed(paragraph.runs):
        paragraph._element.remove(run._element)


def try_insert_image(paragraph: Paragraph, full_text: str, image_paths: dict) -> bool:
    """Insert an image if the placeholder is found in the paragraph."""
    for img_placeholder, img_info in image_paths.items():
        token = f"{{{{{img_placeholder}}}}}"
        if token not in full_text:
            continue

        clear_paragraph_runs(paragraph)

        container = img_info.get("container")
        blob_name = img_info.get("blobName")
        resize_pct = img_info.get("resizePercentage")
        jpeg_quality = img_info.get("jpegQuality")
        width_cm = img_info.get("widthCm")

        final_img_source = get_image(container, blob_name, resize_pct, jpeg_quality)

        if not final_img_source:
            logger.warning(
                "Missing image for placeholder '%s' (container=%r, blobName=%r); using fallback",
                img_placeholder,
                container,
                blob_name,
            )
            final_img_source = get_unavailable_image()

        if final_img_source:
            target_width = width_cm
            run = paragraph.add_run()
            run.add_picture(final_img_source, width=Cm(target_width))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return True

        logger.error(
            "Unable to insert fallback image for placeholder '%s'",
            img_placeholder,
        )
        return False

    return False


def apply_text_replacements(paragraph: Paragraph, replacements: dict) -> None:
    """Replace tokens in a paragraph with their values."""
    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text

    for placeholder, value in replacements.items():
        new_text = new_text.replace(placeholder, str(value))

    if new_text != full_text:
        clear_paragraph_runs(paragraph)
        add_colored_text(paragraph, new_text)


def iter_paragraphs(element) -> list[Paragraph]:
    """Return a flat list of paragraphs from a paragraph or a table."""
    if isinstance(element, Paragraph):
        return [element]

    if isinstance(element, Table):
        result: list[Paragraph] = []
        for row in element.rows:
            for cell in row.cells:
                result.extend(cell.paragraphs)
        return result

    return []


def replace_in_element(element, replacements: dict, image_paths: dict) -> None:
    """Apply image and text replacements on each paragraph inside an element."""
    for paragraph in iter_paragraphs(element):
        full_text = "".join(run.text for run in paragraph.runs)

        image_inserted = try_insert_image(paragraph, full_text, image_paths)
        if not image_inserted:
            apply_text_replacements(paragraph, replacements)


def generate_verification_report_bytes(
    template: bytes,
    replacements: dict,
    image_paths: dict,
) -> io.BytesIO | None:
    """Load a DOCX template, apply placeholders and return a new DOCX stream."""
    try:
        buf = io.BytesIO(template)
        document = Document(buf)
    except Exception:
        logger.exception("Error loading DOCX template")
        return None

    for paragraph in document.paragraphs:
        replace_in_element(paragraph, replacements, image_paths)

    for table in document.tables:
        replace_in_element(table, replacements, image_paths)

    out_stream = io.BytesIO()
    document.save(out_stream)
    out_stream.seek(0)
    return out_stream
